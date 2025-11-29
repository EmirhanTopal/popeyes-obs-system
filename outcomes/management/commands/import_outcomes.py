import os
import docx
import re
from django.core.management.base import BaseCommand
from outcomes.models import ProgramOutcome


class OutcomeImporter:
    @staticmethod
    def extract_number_from_text(text):
        """Metinden outcome numarasını çıkar"""
        match = re.match(r'^(\d+)[\s\-\.]', text.strip())
        return int(match.group(1)) if match else None

    @staticmethod
    def clean_description(text):
        """Outcome açıklamasını temizle"""
        cleaned = re.sub(r'^\d+[\s\-\.]*', '', text.strip())
        cleaned = re.sub(r'\s+', ' ', cleaned)
        return cleaned.strip()

    @staticmethod
    def detect_program_type(file_path):
        """
        Dosya içeriğine göre program tipini otomatik algıla.
        .docx dosyasının ilk 5 paragrafına bakarak bölüm kodunu belirler.
        """
        try:
            doc = docx.Document(file_path)
            content = " ".join([p.text for p in doc.paragraphs[:5]]).lower()

            # === Anahtar kelime eşleştirme ===
            if "biyomedikal" in content or "biomedical" in content:
                return "BME"  # 🔹 SENİN DEPARTMENT CODE'UN
            elif "bilgisayar" in content or "computer" in content:
                return "CSE"
            elif "elektrik" in content or "elektronik" in content:
                return "CSE2"
            else:
                # Basit keyword skoru
                biomedical_keywords = ['tıp', 'medikal', 'sağlık', 'cihaz', 'biyoloji']
                computer_keywords = ['yazılım', 'donanım', 'programlama', 'bilişim']

                biomedical_score = sum(k in content for k in biomedical_keywords)
                computer_score = sum(k in content for k in computer_keywords)

                if biomedical_score > computer_score:
                    return "BME"
                else:
                    return "CSE"

        except Exception as e:
            print(f"Program tipi algılanırken hata: {e}")
            return None

    @staticmethod
    def parse_docx_file(file_path, program_type=None):
        """Docx dosyasını parse et - outcome numaralarını ve açıklamaları çıkar."""
        if not program_type:
            program_type = OutcomeImporter.detect_program_type(file_path)

        outcomes = []
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and any(ch.isdigit() for ch in text[:10]):
                    number = OutcomeImporter.extract_number_from_text(text)
                    if number:
                        description = OutcomeImporter.clean_description(text)
                        outcomes.append({
                            "number": number,
                            "description": description
                        })

            return outcomes, program_type
        except Exception as e:
            print(f"Dosya okunurken hata: {e}")
            return [], None


class Command(BaseCommand):
    help = "YÖK Program Outcomes dokümanlarını otomatik içe aktarır."

    def add_arguments(self, parser):
        parser.add_argument("--files", nargs="+", type=str, help="Docx dosya yolları")
        parser.add_argument("--auto-detect", action="store_true", help="Program tipini otomatik algıla")

    def handle(self, *args, **options):
        importer = OutcomeImporter()
        processed_files = 0

        if options["files"]:
            for file_path in options["files"]:
                if not os.path.exists(file_path):
                    self.stdout.write(self.style.ERROR(f"❌ Dosya bulunamadı: {file_path}"))
                    continue

                program_type = None
                if options["auto_detect"]:
                    program_type = importer.detect_program_type(file_path)
                    self.stdout.write(f"{file_path} → {program_type} olarak algılandı")

                self.import_single_file(importer, file_path, program_type)
                processed_files += 1

        if processed_files == 0:
            self.stdout.write(self.style.WARNING("⚠️ Hiç dosya işlenmedi."))
        else:
            self.stdout.write(self.style.SUCCESS(f"✅ {processed_files} dosya başarıyla işlendi!"))

    def import_single_file(self, importer, file_path, program_type=None):
        """Tek bir docx dosyasını import eder."""
        try:
            outcomes, detected_program = importer.parse_docx_file(file_path, program_type)

            if not detected_program:
                self.stdout.write(self.style.ERROR(f"Program tipi algılanamadı: {file_path}"))
                return

            self.stdout.write(f"{file_path} → {detected_program} outcomes import ediliyor...")

            for outcome in outcomes:
                ProgramOutcome.objects.update_or_create(
                    program=detected_program,
                    outcome_number=outcome["number"],
                    defaults={"description": outcome["description"]},
                )
                self.stdout.write(
                    self.style.SUCCESS(f"{detected_program} Outcome {outcome['number']} eklendi/güncellendi")
                )

        except Exception as e:
            self.stdout.write(self.style.ERROR(f"{file_path} işlenirken hata: {e}"))
